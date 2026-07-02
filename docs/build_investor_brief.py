from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Intern Nexus Investor Brief.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 101, 116)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALL_OUT = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "B8C4D2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def set_run(run, *, size=None, color=None, bold=None, italic=None, font="Calibri") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def para(doc, text="", *, style=None, size=11, color=None, bold=False, italic=False, before=0, after=6, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run(r, size=size, color=color, bold=bold, italic=italic)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run(r, size=10.7)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run(r, size=10.7)
    return p


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=16, color=BLUE, bold=True)
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, size=13, color=BLUE, bold=True)
    return p


def h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=12, color=DARK_BLUE, bold=True)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_borders(table, color="D6DEE8", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALL_OUT)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run(r, size=11.2, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run(r2, size=10.5, color=RGBColor(40, 48, 58))
    para(doc, "", after=3)


def add_matrix(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    set_col_widths(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run(r, size=9.7, color=NAVY, bold=True)
    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            set_cell_margins(row[i])
            row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = row[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(text))
            set_run(r, size=9.5, color=RGBColor(24, 31, 42), bold=(i == 0))
    para(doc, "", after=5)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["List Bullet", "List Number", "List Bullet 2"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.7)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Intern Nexus | Investor Brief")
    set_run(run, size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Prepared for investor discussion")
    set_run(run, size=8.5, color=MUTED)


def cover(doc):
    para(doc, "Investor Brief", size=10.5, color=BLUE, bold=True, after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Intern Nexus")
    set_run(r, size=30, color=NAVY, bold=True)
    p = para(
        doc,
        "Real-time internship management and supervision platform",
        size=14,
        color=MUTED,
        after=18,
    )
    add_matrix(
        doc,
        ["Prepared For", "Current Stage", "Technical Status", "Deployment Direction"],
        [
            [
                "Prospective investors and partners",
                "Localhost prototype with core workflows working",
                "Security hardening, role-based access, backups, cleanup, and audit logs added",
                "PostgreSQL-backed production deployment",
            ]
        ],
        [1.7, 1.6, 2.0, 1.2],
        header_fill=LIGHT_BLUE,
    )
    add_callout(
        doc,
        "Positioning statement",
        "Intern Nexus gives universities, companies, and interns one controlled system for placement visibility, GPS attendance, daily reporting, complaints, evaluations, document sharing, and audit-ready supervision.",
    )
    para(
        doc,
        "This document is designed as a presentation companion: it explains what the system does, how it works technically, what has already been built, and what must be added before a production investor pilot.",
        size=10.8,
        color=RGBColor(41, 51, 63),
        after=8,
    )
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_document(doc)
    cover(doc)

    h1(doc, "1. Executive Summary")
    add_callout(
        doc,
        "The problem",
        "Internship programs are often managed through phone calls, spreadsheets, WhatsApp messages, paper attendance, and scattered supervisor reports. That creates weak visibility, late intervention, attendance disputes, and limited accountability.",
    )
    add_callout(
        doc,
        "The solution",
        "Intern Nexus centralizes the internship lifecycle. Each user logs into a role-specific dashboard, performs only the actions allowed for that role, and leaves a traceable record for attendance, reports, complaints, documents, evaluations, and administrative decisions.",
    )
    bullet(doc, "For universities: better visibility over student placements, attendance, reports, complaints, and supervisor follow-up.")
    bullet(doc, "For host companies: easier task assignment, report review, attendance confirmation, and intern evaluation.")
    bullet(doc, "For interns: one place to check in, submit reports, receive tasks, access documents, and raise complaints.")
    bullet(doc, "For administrators: centralized user approval, placement management, access control, backup/restore, stale-data cleanup, and audit tracking.")

    h1(doc, "2. How The System Works")
    add_matrix(
        doc,
        ["Role", "Main Dashboard Purpose", "Key Permissions"],
        [
            ["Administrator", "Full system control center", "Approve accounts, manage placements, view all records, reset passwords, backup/restore, cleanup data, audit system actions."],
            ["Intern", "Personal internship workspace", "Check in/out with GPS, submit reports, view own tasks, download documents, submit complaints, view own evaluations."],
            ["Company Supervisor", "Company supervision queue", "Review assigned interns, assign tasks, evaluate interns manually, approve reports before university review, respond to assigned complaints."],
            ["University Supervisor", "Academic supervision dashboard", "Review reports after company approval, monitor assigned students, receive attendance notifications, handle escalated complaints."],
        ],
        [1.35, 2.15, 3.0],
        header_fill=LIGHT_BLUE,
    )
    h2(doc, "Core Workflow")
    numbered(doc, "A user creates an account request with their role, assigned ID, organization, and password.")
    numbered(doc, "The administrator approves or rejects the request in the Security Center.")
    numbered(doc, "Interns are assigned to companies and universities through placement records.")
    numbered(doc, "Interns check in/out using browser GPS, which is compared against workplace coordinates and radius.")
    numbered(doc, "Interns submit structured weekly reports to the company supervisor.")
    numbered(doc, "The company supervisor approves the report and forwards it to the university supervisor.")
    numbered(doc, "The university supervisor gives final approval, creating an auditable supervision chain.")
    numbered(doc, "Complaints, evaluations, shared documents, notifications, and audit logs support the supervision lifecycle.")

    h1(doc, "3. Current Product Modules")
    add_matrix(
        doc,
        ["Module", "What It Does", "Investor Value"],
        [
            ["Authentication and RBAC", "Approved users receive signed access tokens and see only their allowed dashboard.", "Protects sensitive student, company, and supervisor data."],
            ["Placement Directory", "Tracks intern, student number, company, university, department, supervisor, coordinates, and radius.", "Creates operational visibility over where every intern is working."],
            ["GPS Attendance", "Uses live browser GPS and geofence rules to accept or reject check-in/out events.", "Reduces false attendance and improves accountability."],
            ["Reports Workflow", "Intern reports move from company approval to university approval.", "Supports formal internship assessment and audit trails."],
            ["Complaints and Comments", "Routes issues according to user role and complaint category.", "Creates controlled escalation and response handling."],
            ["Documents", "Supervisors/admins upload real PDF/DOC/DOCX files for approved users to download.", "Turns the system into a document distribution hub."],
            ["Evaluations", "Company supervisors manually evaluate interns while attendance remains system-managed.", "Combines human judgment with automated evidence."],
            ["Security Center", "Account approvals, backup/restore, password reset, cleanup, and audit stream.", "Gives administrators operational control before deployment."],
        ],
        [1.45, 2.75, 2.3],
    )

    h1(doc, "4. Technology Stack")
    h2(doc, "Frontend")
    add_matrix(
        doc,
        ["Layer", "Technology Used", "Why It Was Chosen"],
        [
            ["Application Framework", "React with TypeScript", "Component-based UI, type safety, role-specific dashboards, and maintainable frontend logic."],
            ["Build Tool", "Vite", "Fast local development, modern bundling, and a clean path to production builds."],
            ["UI Styling", "CSS with responsive layouts", "Custom polished interface without heavy template dependency."],
            ["Icons", "Lucide React", "Clean icons for navigation, forms, actions, and dashboard controls."],
            ["Data Fetching", "Fetch API and TanStack Query health polling", "Simple REST integration with room to scale into richer caching later."],
        ],
        [1.5, 2.15, 2.85],
        header_fill=LIGHT_BLUE,
    )
    h2(doc, "Backend")
    add_matrix(
        doc,
        ["Layer", "Technology Used", "Purpose"],
        [
            ["Runtime", "Node.js ESM HTTP server", "Hosts the REST API used by all dashboards and workflow actions."],
            ["Authentication", "Signed bearer tokens", "Protects private API routes and prevents role spoofing from the frontend."],
            ["Password Security", "bcryptjs hashing", "Stops plaintext password storage and supports migrated local accounts."],
            ["Authorization", "Server-side role-based access control", "Enforces admin, intern, company supervisor, and university supervisor permissions."],
            ["Attendance Logic", "Haversine distance calculation and strict geofence mode", "Compares live GPS coordinates against assigned workplace coordinates."],
            ["File Handling", "Local document storage with file type, MIME, size, and basic signature validation", "Allows real document upload/download during localhost stage."],
            ["Audit Logging", "Append-only local audit events", "Records sensitive actions such as login, cleanup, backup, reports, attendance, and password reset."],
        ],
        [1.45, 2.35, 2.7],
    )

    h1(doc, "5. Database Plan For Deployment")
    add_callout(
        doc,
        "Recommendation",
        "Use PostgreSQL for production deployment. The current JSON store is acceptable for localhost development only, but PostgreSQL is the right database for real users, relationships, transactions, reporting, backups, and scale.",
    )
    add_matrix(
        doc,
        ["Current Local Storage", "Deployment Database", "File Storage"],
        [
            [
                "server/data-store.json for fast localhost testing, backup/restore, and prototyping.",
                "PostgreSQL hosted on Supabase, Neon, Railway, Render, AWS RDS, or another managed Postgres provider.",
                "Object storage for uploaded documents, such as Supabase Storage, S3-compatible storage, or cloud provider storage.",
            ]
        ],
        [2.1, 2.45, 1.95],
        header_fill=LIGHT_BLUE,
    )
    h2(doc, "Suggested Production Tables")
    bullet(doc, "users, account_requests, roles, companies, universities, placements")
    bullet(doc, "attendance_events, reports, report_approvals, evaluations")
    bullet(doc, "complaints, complaint_comments, documents, notifications, audit_logs")
    bullet(doc, "password_reset_tokens, sessions, backup_jobs, integration_settings")
    h2(doc, "Why PostgreSQL Helps Investors Trust The Platform")
    bullet(doc, "It supports relational integrity between interns, supervisors, companies, universities, reports, and attendance.")
    bullet(doc, "It can enforce uniqueness, prevent accidental duplicates, and support clean reporting queries.")
    bullet(doc, "It is deployment-friendly and widely supported by managed cloud providers.")
    bullet(doc, "It gives a clean path to analytics, dashboards, and investor-grade operational reporting.")

    h1(doc, "6. Deployment Readiness: What Has Been Added")
    add_matrix(
        doc,
        ["Area", "Current Improvement", "Status"],
        [
            ["Authentication", "Private APIs require bearer tokens.", "Implemented"],
            ["Password Security", "Passwords migrated from plaintext to bcrypt hashes.", "Implemented"],
            ["Role Enforcement", "Backend derives user role from verified token.", "Implemented"],
            ["Report Approval Security", "Company/university approval is restricted by assignment.", "Implemented"],
            ["Geofence", "Strict geofence is the default; local test mode is explicit.", "Implemented"],
            ["Documents", "Real PDF/DOC/DOCX upload and download with validation.", "Implemented"],
            ["Admin Tools", "Backup, restore, cleanup, password reset, and audit stream.", "Implemented"],
            ["Stale Data Cleanup", "Likely test records removed after backup.", "Completed"],
        ],
        [1.7, 3.5, 1.3],
        header_fill=LIGHT_BLUE,
    )

    h1(doc, "7. What Must Be Added Before Deployment")
    add_matrix(
        doc,
        ["Priority", "Deployment Item", "Reason"],
        [
            ["Critical", "Move from JSON to PostgreSQL", "Prevents data corruption and supports real users, relationships, and reporting."],
            ["Critical", "Configure production secrets and HTTPS hosting", "Protects authentication tokens, passwords, and user data in transit."],
            ["Critical", "Use cloud object storage for documents", "Keeps uploads reliable, backed up, and available across deployments."],
            ["High", "Email-based password reset", "Allows users to recover accounts without administrator manual reset."],
            ["High", "Automated backup schedule", "Protects against accidental deletion, corruption, or failed deployments."],
            ["High", "Monitoring and error logging", "Helps operators detect failed login, upload, attendance, and approval flows."],
            ["High", "Admin audit trail page with filters", "Makes security and operations reviewable by non-technical administrators."],
            ["Medium", "Google sign-in after Google Cloud setup", "Useful for easier onboarding, but only after OAuth credentials are approved."],
            ["Medium", "Email/SMS/push notifications", "Improves engagement and supervisor response times."],
            ["Medium", "Full malware scanning for uploads", "Adds production-grade protection beyond current basic validation."],
        ],
        [1.0, 2.35, 3.15],
    )

    h1(doc, "8. Deployment Architecture")
    add_matrix(
        doc,
        ["Component", "Recommended Deployment Choice", "Notes"],
        [
            ["Frontend", "Vite production build hosted on Vercel, Netlify, Render, or similar", "Static frontend can be deployed separately from the API."],
            ["Backend API", "Node.js service on Render, Railway, Fly.io, DigitalOcean, or VPS", "Runs authentication, workflows, attendance, uploads, and audit logic."],
            ["Database", "Managed PostgreSQL", "Use Supabase, Neon, Railway Postgres, Render Postgres, or AWS RDS."],
            ["File Storage", "Cloud object storage", "Store uploaded reports/templates/documents outside the app server filesystem."],
            ["Email", "SMTP or API provider", "Resend, SendGrid, Mailgun, Gmail SMTP, or organization SMTP for reset links and alerts."],
            ["Security", "HTTPS, environment secrets, access token expiry, backups, monitoring", "Required before inviting real users."],
        ],
        [1.35, 2.65, 2.5],
        header_fill=LIGHT_BLUE,
    )

    h1(doc, "9. Investor Value Proposition")
    h2(doc, "Why This Product Matters")
    bullet(doc, "Universities need visibility over student welfare, attendance, report quality, and workplace issues.")
    bullet(doc, "Companies need a simple way to manage interns without creating manual spreadsheets.")
    bullet(doc, "Interns need one trusted channel for proof of attendance, reports, documents, and complaints.")
    bullet(doc, "Administrators need reliable controls, audit logs, backups, and permission enforcement.")
    h2(doc, "Why The Timing Is Strong")
    bullet(doc, "Internship programs are growing, but supervision workflows are still fragmented.")
    bullet(doc, "Institutions increasingly need digital evidence, compliance, and transparent reporting.")
    bullet(doc, "The system can expand from one university pilot into multi-institution deployment.")

    h1(doc, "10. Suggested 30/60/90 Day Plan")
    add_matrix(
        doc,
        ["Phase", "Main Work", "Outcome"],
        [
            ["First 30 days", "PostgreSQL migration, deployment environment, cloud storage, email reset, production secrets.", "Secure pilot-ready backend and database."],
            ["Days 31-60", "Pilot with one university and selected companies; collect workflow feedback; refine reports and notifications.", "Validated user workflows and early adoption evidence."],
            ["Days 61-90", "Analytics, admin reporting, audit page improvements, institution onboarding, payment/pricing exploration.", "Investor-ready traction story and scalable operating model."],
        ],
        [1.2, 3.3, 2.0],
        header_fill=LIGHT_BLUE,
    )

    h1(doc, "11. Investor Discussion Talking Points")
    bullet(doc, "Intern Nexus is not only an attendance tool; it is a supervision operating system for internships.")
    bullet(doc, "The product already has role dashboards, GPS attendance, report approval flow, complaints, documents, evaluations, audit logs, backup/restore, and admin security controls.")
    bullet(doc, "The next major technical milestone is PostgreSQL deployment, followed by email reset, hosted file storage, and pilot onboarding.")
    bullet(doc, "The first commercial opportunity is a university-led pilot with company supervisors and interns onboarded under one controlled system.")
    bullet(doc, "The long-term opportunity is a multi-institution internship management platform with analytics, compliance reporting, and verified work records.")

    h1(doc, "Appendix: Current Status Snapshot")
    add_matrix(
        doc,
        ["Category", "Current State"],
        [
            ["Local app URL", "Frontend runs on http://127.0.0.1:5173 during development."],
            ["Local API URL", "Backend API runs on http://127.0.0.1:4500 during development."],
            ["Current database", "Local JSON store after cleanup; suitable only for localhost testing."],
            ["Production database recommendation", "PostgreSQL with managed hosting."],
            ["Google sign-in", "Disabled until Google Cloud OAuth credentials are created."],
            ["Security posture", "Bearer-token auth, hashed passwords, RBAC, strict geofence, audit logs, and protected admin tools are implemented."],
        ],
        [2.0, 4.5],
    )

    doc.core_properties.title = "Intern Nexus Investor Brief"
    doc.core_properties.subject = "Investor presentation companion for Intern Nexus"
    doc.core_properties.author = "Intern Nexus Team"
    doc.core_properties.comments = "Generated as an investor-facing technical and deployment readiness brief."
    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT)
